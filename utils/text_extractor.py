"""Extract plain text and metadata from book files (PDF, EPUB, DOCX, FB2, TXT, RTF)."""

import re
from pathlib import Path
from typing import NamedTuple


class ExtractResult(NamedTuple):
    text: str
    title: str
    author: str
    year: str
    publisher: str
    description: str = ""  # summary/annotation from file metadata


_EMPTY = ExtractResult("", "", "", "", "")
_MAX_CHARS = 5000
_DEEP_MAX_CHARS = 8000
_DEEP_MIN_PAGE_CHARS = 80   # pages shorter than this are skipped (title/blank pages)


def extract(filepath: str) -> ExtractResult:
    ext = Path(filepath).suffix.lower()
    try:
        if ext == ".pdf":
            return _pdf(filepath)
        elif ext == ".epub":
            return _epub(filepath)
        elif ext == ".docx":
            return _docx(filepath)
        elif ext == ".doc":
            return _doc(filepath)
        elif ext == ".fb2":
            return _fb2(filepath)
        elif ext == ".txt":
            return _txt(filepath)
        elif ext == ".rtf":
            return _rtf(filepath)
        elif ext == ".zip":
            return _zip(filepath)
        else:
            return _EMPTY
    except Exception:
        return _EMPTY


# ---------------------------------------------------------------- extractors

def _pdf(filepath: str) -> ExtractResult:
    import fitz  # PyMuPDF
    doc = fitz.open(filepath)
    meta = doc.metadata or {}
    title = (meta.get("title") or "").strip()
    author = (meta.get("author") or "").strip()
    description = (meta.get("subject") or "").strip()
    text = ""
    for i in range(min(5, doc.page_count)):
        text += doc[i].get_text()
        if len(text) >= _MAX_CHARS:
            break
    doc.close()
    return ExtractResult(text[:_MAX_CHARS], title, author, "", "", description)


def _epub(filepath: str) -> ExtractResult:
    import warnings
    warnings.filterwarnings("ignore")
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup

    book = epub.read_epub(filepath, options={"ignore_ncx": True})

    titles = book.get_metadata("DC", "title")
    title = titles[0][0] if titles else ""

    creators = book.get_metadata("DC", "creator")
    author = "; ".join(c[0] for c in creators) if creators else ""

    descs = book.get_metadata("DC", "description")
    description = descs[0][0].strip() if descs else ""

    text = ""
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), "html.parser")
        text += soup.get_text(separator=" ", strip=True) + " "
        if len(text) >= _MAX_CHARS:
            break

    return ExtractResult(text[:_MAX_CHARS], title, author, "", "", description)


def _docx(filepath: str) -> ExtractResult:
    from docx import Document
    doc = Document(filepath)
    props = doc.core_properties
    title = (props.title or "").strip()
    author = (props.author or "").strip()
    description = (props.description or "").strip()
    text = "\n".join(p.text for p in doc.paragraphs[:150] if p.text.strip())
    return ExtractResult(text[:_MAX_CHARS], title, author, "", "", description)


def _doc(filepath: str) -> ExtractResult:
    try:
        return _docx(filepath)  # works for some .doc files
    except Exception:
        return _EMPTY


def _fb2(filepath: str) -> ExtractResult:
    import xml.etree.ElementTree as ET
    NS = "http://www.gribuser.ru/xml/fictionbook/2.0"
    ns = {"fb": NS}
    tree = ET.parse(filepath)
    root = tree.getroot()

    title = author = year = publisher = ""

    annotation = ""
    desc = root.find("fb:description", ns)
    if desc is not None:
        ti = desc.find("fb:title-info", ns)
        if ti is not None:
            bt = ti.find("fb:book-title", ns)
            if bt is not None and bt.text:
                title = bt.text.strip()
            auth = ti.find("fb:author", ns)
            if auth is not None:
                parts = [
                    (auth.find(f"fb:{t}", ns) or type("", (), {"text": ""})()).text or ""
                    for t in ("first-name", "middle-name", "last-name")
                ]
                author = " ".join(p.strip() for p in parts if p.strip())
            annot = ti.find("fb:annotation", ns)
            if annot is not None:
                annotation = " ".join(annot.itertext()).strip()
        pi = desc.find("fb:publish-info", ns)
        if pi is not None:
            for tag, var in (("fb:year", "year"), ("fb:publisher", "publisher")):
                el = pi.find(tag, ns)
                if el is not None and el.text:
                    locals()[var]  # satisfy linter; actual assignment below
            yr = pi.find("fb:year", ns)
            if yr is not None and yr.text:
                year = yr.text.strip()
            pub = pi.find("fb:publisher", ns)
            if pub is not None and pub.text:
                publisher = pub.text.strip()

    body = root.find("fb:body", ns)
    text = (" ".join(body.itertext())[:_MAX_CHARS]) if body is not None else ""
    return ExtractResult(text, title, author, year, publisher, annotation)


def _txt(filepath: str) -> ExtractResult:
    for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            with open(filepath, "r", encoding=enc) as f:
                return ExtractResult(f.read(_MAX_CHARS), "", "", "", "")
        except UnicodeDecodeError:
            continue
    return _EMPTY


def _rtf(filepath: str) -> ExtractResult:
    try:
        with open(filepath, "rb") as f:
            raw = f.read(10000).decode("latin-1", errors="ignore")
        text = re.sub(r'\\[a-z]+[-\d]*\s?', ' ', raw)
        text = re.sub(r'[{}\\]', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        return ExtractResult(text[:_MAX_CHARS], "", "", "", "")
    except Exception:
        return _EMPTY


def _pdf_pages(doc, max_pages: int, max_chars: int) -> str:
    """Read up to max_pages non-blank pages from an open fitz document."""
    text = ""
    collected = 0
    for i in range(min(max_pages * 2, doc.page_count)):
        if collected >= max_pages or len(text) >= max_chars:
            break
        page_text = doc[i].get_text().strip()
        if len(page_text) >= _DEEP_MIN_PAGE_CHARS:
            text += page_text + "\n"
            collected += 1
    return text[:max_chars]


def _pdf_deep(filepath: str, max_pages: int = 15) -> ExtractResult:
    import fitz
    doc = fitz.open(filepath)
    meta = doc.metadata or {}
    title = (meta.get("title") or "").strip()
    author = (meta.get("author") or "").strip()
    description = (meta.get("subject") or "").strip()
    text = _pdf_pages(doc, max_pages, _DEEP_MAX_CHARS)
    doc.close()
    return ExtractResult(text, title, author, "", "", description)


def _zip_deep(filepath: str, max_pages: int = 15) -> ExtractResult:
    import zipfile
    try:
        with zipfile.ZipFile(filepath, "r") as zf:
            pdf_entries = [n for n in zf.namelist()
                           if n.lower().endswith(".pdf") and not n.endswith("/")]
            if len(pdf_entries) == 1:
                import fitz
                data = zf.read(pdf_entries[0])
                doc = fitz.open(stream=data, filetype="pdf")
                meta = doc.metadata or {}
                title = (meta.get("title") or "").strip()
                author = (meta.get("author") or "").strip()
                description = (meta.get("subject") or "").strip()
                text = _pdf_pages(doc, max_pages, _DEEP_MAX_CHARS)
                doc.close()
                return ExtractResult(text, title, author, "", "", description)
    except Exception:
        pass
    return _EMPTY


def extract_deep(filepath: str, max_pages: int = 15) -> ExtractResult:
    """Like extract() but reads up to max_pages non-blank pages (for summaries)."""
    ext = Path(filepath).suffix.lower()
    try:
        if ext == ".pdf":
            return _pdf_deep(filepath, max_pages)
        elif ext == ".zip":
            return _zip_deep(filepath, max_pages)
        else:
            return extract(filepath)
    except Exception:
        return _EMPTY


def _pdf_from_bytes(data: bytes) -> ExtractResult:
    import fitz
    doc = fitz.open(stream=data, filetype="pdf")
    meta = doc.metadata or {}
    title = (meta.get("title") or "").strip()
    author = (meta.get("author") or "").strip()
    description = (meta.get("subject") or "").strip()
    text = ""
    for i in range(min(5, doc.page_count)):
        text += doc[i].get_text()
        if len(text) >= _MAX_CHARS:
            break
    doc.close()
    return ExtractResult(text[:_MAX_CHARS], title, author, "", "", description)


def _zip(filepath: str) -> ExtractResult:
    import zipfile
    try:
        with zipfile.ZipFile(filepath, "r") as zf:
            pdf_entries = [n for n in zf.namelist()
                           if n.lower().endswith(".pdf") and not n.endswith("/")]
            if len(pdf_entries) == 1:
                return _pdf_from_bytes(zf.read(pdf_entries[0]))
    except Exception:
        pass
    return _EMPTY


def extract_description(filepath: str) -> str:
    """Read only the description/annotation from file metadata (fast, no body text)."""
    ext = Path(filepath).suffix.lower()
    try:
        if ext == ".zip":
            import zipfile
            with zipfile.ZipFile(filepath, "r") as zf:
                pdf_entries = [n for n in zf.namelist()
                               if n.lower().endswith(".pdf") and not n.endswith("/")]
                if len(pdf_entries) == 1:
                    import fitz
                    doc = fitz.open(stream=zf.read(pdf_entries[0]), filetype="pdf")
                    meta = doc.metadata or {}
                    doc.close()
                    return (meta.get("subject") or "").strip()
            return ""
        if ext == ".pdf":
            import fitz
            doc = fitz.open(filepath)
            meta = doc.metadata or {}
            doc.close()
            return (meta.get("subject") or "").strip()
        elif ext == ".epub":
            import warnings
            warnings.filterwarnings("ignore")
            from ebooklib import epub
            book = epub.read_epub(filepath, options={"ignore_ncx": True})
            descs = book.get_metadata("DC", "description")
            return descs[0][0].strip() if descs else ""
        elif ext in (".docx", ".doc"):
            from docx import Document
            doc = Document(filepath)
            return (doc.core_properties.description or "").strip()
        elif ext == ".fb2":
            import xml.etree.ElementTree as ET
            NS = "http://www.gribuser.ru/xml/fictionbook/2.0"
            ns = {"fb": NS}
            root = ET.parse(filepath).getroot()
            desc = root.find("fb:description", ns)
            if desc is not None:
                ti = desc.find("fb:title-info", ns)
                if ti is not None:
                    annot = ti.find("fb:annotation", ns)
                    if annot is not None:
                        return " ".join(annot.itertext()).strip()
    except Exception:
        pass
    return ""
